from docx import Document
import random
import mysql.connector
from datetime import datetime

# Predefined random values for hobbies, extracurricular activities, and soft skills
RANDOM_HOBBIES = ["Reading", "Traveling", "Photography", "Cooking", "Gardening"]
RANDOM_EXTRACURRICULAR = [
    "Volunteering at NGOs",
    "Participated in Hackathons",
    "Member of Debate Club",
    "Sports Activities",
    "Organized Cultural Events",
]
RANDOM_SOFT_SKILLS = ["Communication", "Teamwork", "Problem-Solving", "Adaptability", "Leadership"]

# Function to connect to MySQL
def connect_to_db():
    return mysql.connector.connect(
        host="localhost",  # Change to your MySQL host
        user="root",       # Your MySQL username
        password="Neha@123",  # Your MySQL password
        database="resumeDB"   # The database you created
    )

# Function to insert data into MySQL
def insert_resume_data(document_name, creation_time):
    conn = connect_to_db()
    cursor = conn.cursor()
    
    query = "INSERT INTO resumes (document_name, creation_time) VALUES (%s, %s)"
    cursor.execute(query, (document_name, creation_time))
    
    conn.commit()  # Commit the transaction
    cursor.close()
    conn.close()


# Function to collect user details
def collect_user_details():
    user_details = {}
    print("----------------------------WELCOME TO RESUME BUILDER------------------------------")
    print()
    print("-------------------------------Enter your basic details----------------------------")
    print()
    # Basic Details (compulsory)
    while True:
        user_details['name'] = input("Name: ").strip()
        if user_details['name']:
            break
        else:
            print("Name is required. Please enter your name.")

    while True:
        user_details['address'] = input("Address: ").strip()
        if user_details['address']:
            break
        else:
            print("Address is required. Please enter your address.")

    while True:
        user_details['state'] = input("State: ").strip()
        if user_details['state']:
            break
        else:
            print("State is required. Please enter your state.")

    while True:
        user_details['pincode'] = input("Pincode: ").strip()
        if user_details['pincode']:
            break
        else:
            print("Pincode is required. Please enter your pincode.")

    while True:
        user_details['country'] = input("Country: ").strip()
        if user_details['country']:
            break
        else:
            print("Country is required. Please enter your country.")

    while True:
        user_details['email'] = input("Email: ").strip()
        if user_details['email']:
            break
        else:
            print("Email is required. Please enter your email.")

    while True:
        user_details['phone'] = input("Phone: ").strip()
        if user_details['phone']:
            break
        else:
            print("Phone is required. Please enter your phone.")

    # Academic Details (compulsory)
    print("------------------------------Enter your academic details--------------------------")
    print()
    while True:
        user_details['course'] = input("Graducation Course: ").strip()
        if user_details['course']:
            break
        else:
            print("Course is required. Please enter your course.")

    while True:
        user_details['university'] = input("University Name: ").strip()
        if user_details['university']:
            break
        else:
            print("University Name is required. Please enter your university name.")

    while True:
        user_details['grad-year'] = input("Graduation Year : ").strip()
        if user_details['grad-year']:
            break
        else:
            print("Graduation Year is required. Please enter your graduation year.")

    while True:
        user_details['status'] = input("Status (Ongoing or completed): ").strip()
        if user_details['status']:
            break
        else:
            print("status is required. Please enter your status.")


    while True:
        user_details['cgpa'] = input("CGPA: ").strip()
        if user_details['cgpa']:
            break
        else:
            print("CGPA is required. Please enter your CGPA.")

    while True:
        user_details['tot'] = input("Total CGPA: ").strip()
        if user_details['tot']:
            break
        else:
            print("Total CGPA is required. Please enter your total CGPA.")

    while True:
        user_details['board-name-12'] = input("Board Name (Class 12): ").strip()
        if user_details['board-name-12']:
            break
        else:
            print("Board Name (Class 12) is required. Please enter the board name.")

    while True:
        user_details['school-name-12'] = input("School Name (Class 12): ").strip()
        if user_details['school-name-12']:
            break
        else:
            print("School Name (Class 12) is required. Please enter the school name.")

    while True:
        user_details['year-12'] = input("Year of Passing (Class 12): ").strip()
        if user_details['year-12']:
            break
        else:
            print("Year of Passing (Class 12) is required. Please enter the year.")

    while True:
        user_details['perc-12'] = input("Percentage (Class 12): ").strip()
        if user_details['perc-12']:
            break
        else:
            print("Percentage (Class 12) is required. Please enter the percentage.")

    while True:
        user_details['board-name-10'] = input("Board Name (Class 10): ").strip()
        if user_details['board-name-10']:
            break
        else:
            print("Board Name (Class 10) is required. Please enter the board name.")

    while True:
        user_details['school-name-10'] = input("School Name (Class 10): ").strip()
        if user_details['school-name-10']:
            break
        else:
            print("School Name (Class 10) is required. Please enter the school name.")

    while True:
        user_details['year-10'] = input("Year of Passing (Class 10): ").strip()
        if user_details['year-10']:
            break
        else:
            print("Year of Passing (Class 10) is required. Please enter the year.")

    while True:
        user_details['perc-10'] = input("Percentage (Class 10): ").strip()
        if user_details['perc-10']:
            break
        else:
            print("Percentage (Class 10) is required. Please enter the percentage.")

    # You can continue with the rest of the fields in a similar way

    # Internship details (ask user first if they want to add)
    print("-----------------------------Enter your Internship details-------------------------")
    print()
    has_internship = input("Do you want to add internship details? (yes/no): ").strip().lower()

    if has_internship == "yes":
        user_details['intern-position'] = input("Internship Position: ").strip()
        user_details['company-name'] = input("Company Name: ").strip()
        user_details['duration'] = input("Duration (e.g., 3 months): ").strip()
        user_details['internship-content'] = input("Internship Description: ").strip()
    else:
        user_details['intern-position'] = "Not Applicable"
        user_details['company-name'] = ""
        user_details['duration'] = ""
        user_details['internship-content'] = "No internship details provided."


    # technical Skills
    print("------------------------------Enter your technical skills--------------------------")
    print()
    user_details['os'] = input("Operating System you are familiar with(comma-separated): ")
    user_details['skills'] = input("Technical Skills (comma-separated): ")
    user_details['software'] = input("Software Skills:(comma-separated): ")

    # Project details
    print("------------------------------Enter your Project details---------------------------")
    print()
    user_details['projectname'] = input("Project Name: ")
    user_details['project-year'] = input("Project Year: ")
    user_details['project-content'] = input("Project Description: ")

    # Certifications
    print("---------------------------Enter your certification details------------------------")
    print()
    user_details['course-name'] = input("Certification Name: ")
    user_details['platform'] = input("Certification Platform: ")
    user_details['month'] = input("Certification Month: ")
    user_details['year'] = input("Certification Year: ")

    #hobbies
    print("------------------------------Enter your Hobbies(optional)-------------------------")
    print()
    user_details['h1'] = input("Hobby 1: ") or random.choice(RANDOM_HOBBIES)
    user_details['h2'] = input("Hobby 2: ") or random.choice(RANDOM_HOBBIES)
    user_details['h3'] = input("Hobby 3: ") or random.choice(RANDOM_HOBBIES)

    #extra curricular
    print("------------------------Enter your Extracurricular activities----------------------")
    print()
    user_details['ea1'] = input("Extracurricular Activity 1: ") or random.choice(RANDOM_EXTRACURRICULAR)
    user_details['ea2'] = input("Extracurricular Activity 2: ") or random.choice(RANDOM_EXTRACURRICULAR)

    #soft skills
    print("-------------------------------Enter your Soft skills-----------------------------")
    print()
    user_details['sk1'] = input("Soft Skill 1: ") or random.choice(RANDOM_SOFT_SKILLS)
    user_details['sk2'] = input("Soft Skill 2: ") or random.choice(RANDOM_SOFT_SKILLS)
    user_details['sk3'] = input("Soft Skill 3: ") or random.choice(RANDOM_SOFT_SKILLS)

    return user_details

# Function to populate the template
def populate_template(template_path, output_path, user_details):
    # Load the template
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in user_details.items():
            if f"{{{{{placeholder}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{placeholder}}}}}", value)

    # Replace placeholders in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in user_details.items():
                    if f"{{{{{placeholder}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{placeholder}}}}}", value)

    # Save the updated document
    doc.save(output_path)
    print(f"Resume created successfully as {output_path}")
    print("You are free to do any further modification as per your needs.")
    print("THANK YOU FOR USING")

    # Insert document details into MySQL
    creation_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    insert_resume_data(output_path, creation_time)

# Main Program
if __name__ == "__main__":
    # Collect user details
    user_details = collect_user_details()

    # Specify the template and output paths
    template_path = r"template1.docx"  # Update with your actual template path
    output_path = rf"C:\Users\hp\Documents\{user_details['name']}Resume.docx"  # Save the resume

    # Populate the template
    populate_template(template_path, output_path, user_details)
